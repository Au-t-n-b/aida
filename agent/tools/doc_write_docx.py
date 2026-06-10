"""
DocWriteDocxTool · 生成 Word 文档

将 Markdown 风格文本写入 .docx 文件。
支持标题（# / ## / ###）、加粗（**text**）、无序列表（- item）、普通段落。
适合 Agent 输出报告草稿、摘要、通知函等场景。
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from .base import Tool, ToolError


class DocWriteDocxTool(Tool):
    @property
    def name(self) -> str:
        return "doc_write_docx"

    @property
    def description(self) -> str:
        return (
            "把 Markdown 风格文本写入 Word（.docx）文件。"
            "支持 # 标题、**加粗**、- 无序列表、普通段落。"
            "路径不存在时自动创建父目录；文件已存在时覆盖。"
            "适合生成报告草稿、摘要、通知函等文本型产物。"
        )

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "输出 .docx 文件路径（绝对或相对）",
                    "minLength": 1,
                },
                "content": {
                    "type": "string",
                    "description": "Markdown 风格文本内容",
                    "minLength": 1,
                },
                "title": {
                    "type": "string",
                    "description": "可选：文档标题（作为第一级标题写入，优先于 content 里的 #）",
                },
            },
            "required": ["path", "content"],
        }

    def execute(
        self,
        path: str,
        content: str,
        title: str = "",
        **_: Any,
    ) -> str:
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.oxml.ns import qn
        except ImportError:
            return ToolError("缺少依赖 python-docx，请运行 pip install python-docx")

        p = Path(path)
        if p.suffix.lower() not in (".docx",):
            p = p.with_suffix(".docx")

        try:
            p.parent.mkdir(parents=True, exist_ok=True)
        except Exception as e:
            return ToolError(f"无法创建目录 {p.parent}：{e}")

        doc = Document()

        if title:
            doc.add_heading(title, level=1)

        for raw_line in content.splitlines():
            line = raw_line.rstrip()
            if not line:
                doc.add_paragraph("")
                continue

            # 标题
            h3 = re.match(r"^#{3}\s+(.*)", line)
            h2 = re.match(r"^#{2}\s+(.*)", line)
            h1 = re.match(r"^#\s+(.*)", line)
            if h1:
                doc.add_heading(h1.group(1), level=1)
            elif h2:
                doc.add_heading(h2.group(1), level=2)
            elif h3:
                doc.add_heading(h3.group(1), level=3)
            # 无序列表
            elif re.match(r"^[-*]\s+", line):
                text = re.sub(r"^[-*]\s+", "", line)
                para = doc.add_paragraph(style="List Bullet")
                _add_inline(para, text)
            # 有序列表
            elif re.match(r"^\d+\.\s+", line):
                text = re.sub(r"^\d+\.\s+", "", line)
                para = doc.add_paragraph(style="List Number")
                _add_inline(para, text)
            # 普通段落
            else:
                para = doc.add_paragraph()
                _add_inline(para, line)

        try:
            doc.save(str(p))
        except Exception as e:
            return ToolError(f"保存 docx 失败：{e}")

        size_kb = p.stat().st_size // 1024
        return f"OK: 已写入 {p}（{size_kb} KB，{len(content.splitlines())} 行输入）"


def _add_inline(para: Any, text: str) -> None:
    """解析行内 **bold** 标记，添加 Run。"""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        m = re.match(r"^\*\*([^*]+)\*\*$", part)
        if m:
            run = para.add_run(m.group(1))
            run.bold = True
        else:
            para.add_run(part)
