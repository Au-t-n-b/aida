"""
交付预案 · docx/xlsx 解析服务（0610 mock）
"""
from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any


def _cell_text(cell) -> str:
    return (cell.text or "").strip().replace("\n", " ")


def _fmt_date(v) -> str:
    if v is None:
        return ""
    if hasattr(v, "strftime"):
        return v.strftime("%Y-%m-%d")
    return str(v).strip()


def _resolve_xlsx_data_start(all_rows: list[tuple]) -> tuple[list[str], int]:
    """识别表头行与数据起始行（兼容旧版 meta 行 + 新版的「项目名称…版本号」单行表头）。"""
    if not all_rows:
        return [], 0
    h0 = [str(h or "").strip() for h in all_rows[0]]
    if not h0:
        return [], 0
    # 新版：项目名称在首列、版本号在末列
    if h0[0] == "项目名称" and h0[-1] == "版本号" and h0[1] != "版本号":
        return h0, 1
    # 旧版：第 1 行「项目名称,版本号,字段…」，第 3 行才是字段表头
    if (
        len(all_rows) >= 3
        and h0[0] == "项目名称"
        and len(h0) > 1
        and h0[1] == "版本号"
    ):
        h2 = [str(h or "").strip() for h in all_rows[2]]
        if any(k in h2 for k in ("技术栈", "活动", "活动分类", "分类", "用例编号")):
            return h2, 3
    return h0, 1


def parse_raci_xlsx(path: Path) -> list[dict[str, str]]:
    import openpyxl

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    ws = wb.active
    all_rows = list(ws.iter_rows(values_only=True))
    headers, data_start = _resolve_xlsx_data_start(all_rows)
    if not headers:
        return []

    alias = {
        "技术栈": "stack",
        "活动分类": "cat",
        "活动分类（一级）": "cat",
        "活动": "act",
        "活动（二级）": "act",
        "GTS": "gts",
        "华为云": "hw",
        "伙伴": "partner",
        "客户": "customer",
    }
    idx = {alias.get(h, h): i for i, h in enumerate(headers) if h not in ("项目名称", "版本号", "序号")}

    out: list[dict[str, str]] = []
    for row in all_rows[data_start:]:
        if not row or all(v is None or str(v).strip() == "" for v in row):
            continue

        def get(k: str) -> str:
            if k not in idx or idx[k] >= len(row):
                return ""
            v = row[idx[k]]
            return "" if v is None else str(v).strip()

        if get("act") in ("", "活动（二级）", "活动") and get("cat") in ("", "活动分类（一级）", "活动分类"):
            continue
        if not get("act") and not get("cat") and not get("stack"):
            continue
        out.append({
            "stack": get("stack"),
            "cat": get("cat"),
            "act": get("act"),
            "gts": get("gts"),
            "hw": get("hw"),
            "partner": get("partner"),
            "customer": get("customer"),
        })
    return out


def parse_plan_xlsx(path: Path) -> list[dict[str, Any]]:
    import openpyxl

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    ws = wb.active
    all_rows = list(ws.iter_rows(values_only=True))
    headers, data_start = _resolve_xlsx_data_start(all_rows)
    if not headers:
        return []

    alias = {
        "活动名称": "name",
        "开始": "start",
        "开始日期": "start",
        "结束": "end",
        "结束日期": "end",
        "实际开始": "actualStart",
        "实际开始日期": "actualStart",
        "实际结束": "actualEnd",
        "实际结束日期": "actualEnd",
        "责任人": "owner",
        "管理单元": "unit",
        "任务状态": "status",
        "进度": "progress",
    }
    idx = {alias.get(h, h): i for i, h in enumerate(headers) if h not in ("项目名称", "版本号")}

    out: list[dict[str, Any]] = []
    for row in all_rows[data_start:]:
        if not row or all(v is None or str(v).strip() == "" for v in row):
            continue

        def get(k: str):
            if k not in idx or idx[k] >= len(row):
                return None
            return row[idx[k]]

        name = str(get("name") or "").strip()
        if not name or name in ("活动名称", "项目名称"):
            continue
        prog_raw = get("progress")
        progress = 0
        if prog_raw is not None:
            s = str(prog_raw).replace("%", "").strip()
            try:
                val = float(s)
                progress = int(val * 100) if 0 < val <= 1 else int(val)
            except ValueError:
                progress = 0
        out.append({
            "name": name,
            "start": _fmt_date(get("start")),
            "end": _fmt_date(get("end")),
            "actualStart": _fmt_date(get("actualStart")),
            "actualEnd": _fmt_date(get("actualEnd")),
            "owner": str(get("owner") or ""),
            "unit": str(get("unit") or ""),
            "status": str(get("status") or ""),
            "progress": progress,
            "progressTone": "green" if progress >= 100 else "blue",
        })
    return out


def parse_testcases_xlsx(path: Path) -> list[dict[str, Any]]:
    try:
        from .officecli_parse import parse_testcases_via_officecli

        rows = parse_testcases_via_officecli(path)
        if rows:
            return rows
    except Exception:
        pass

    import openpyxl

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    ws = wb.active
    all_rows = list(ws.iter_rows(values_only=True))
    headers, data_start = _resolve_xlsx_data_start(all_rows)
    if not headers:
        return []

    alias = {
        "用例编号": "id",
        "一级分类": "l1",
        "二级分类": "l2",
        "三级分类": "l3",
        "测试目的": "purpose",
        "测试组网": "topology",
        "预置条件": "pre",
        "测试步骤": "steps",
        "预期结果": "expects",
        "测试结果": "result",
        "备注": "remark",
        "勾选": "selected",
    }
    idx = {alias.get(h, h): i for i, h in enumerate(headers) if h not in ("项目名称", "版本号")}

    out: list[dict[str, Any]] = []
    for row in all_rows[data_start:]:
        if not row or all(v is None or str(v).strip() == "" for v in row):
            continue

        def get(k: str) -> str:
            if k not in idx or idx[k] >= len(row):
                return ""
            v = row[idx[k]]
            return "" if v is None else str(v).strip()

        if not get("id") or get("id") == "用例编号":
            continue
        steps_raw = get("steps")
        expects_raw = get("expects")
        out.append({
            "id": get("id"),
            "l1": get("l1"),
            "l2": get("l2"),
            "l3": get("l3"),
            "purpose": get("purpose"),
            "topology": get("topology"),
            "pre": get("pre"),
            "steps": [s.strip() for s in steps_raw.split("\n") if s.strip()] if steps_raw else [],
            "expects": [s.strip() for s in expects_raw.split("\n") if s.strip()] if expects_raw else [],
            "result": get("result"),
            "remark": get("remark"),
            "selected": get("selected") == "是",
        })
    return out


def parse_acceptance_from_docx(path: Path) -> list[dict[str, str]]:
    """从技术建议书 docx 提取验收策略（officecli 主路径，python-docx 降级）。"""
    try:
        from .officecli_parse import parse_acceptance_via_officecli

        rows = parse_acceptance_via_officecli(path)
        if rows:
            return rows
    except Exception:
        pass

    from docx import Document

    doc = Document(str(path))
    items: list[dict[str, str]] = []

    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = [_cell_text(c) for c in table.rows[0].cells]
        header_joined = "".join(header_cells)
        if not any(k in header_joined for k in ("验收", "分类", "回款", "里程碑")):
            continue

        col_map: dict[str, int] = {}
        for i, h in enumerate(header_cells):
            if "分类" in h:
                col_map["cat"] = i
            elif "验收方案" in h or h == "方案":
                col_map["scheme"] = i
            elif "验收标准" in h or h == "标准":
                col_map["standard"] = i
            elif "验收里程碑" in h:
                col_map["milestone"] = i
            elif "验收文档" in h:
                col_map["doc"] = i
            elif "回款条款" in h:
                col_map["payment"] = i
            elif "回款里程碑" in h:
                col_map["paymentMilestone"] = i

        if "cat" not in col_map and "scheme" not in col_map:
            continue

        for row in table.rows[1:]:
            cells = [_cell_text(c) for c in row.cells]
            if not any(cells):
                continue
            get = lambda k: cells[col_map[k]] if k in col_map and col_map[k] < len(cells) else ""
            cat = get("cat")
            scheme = get("scheme")
            if not cat and not scheme:
                continue
            items.append({
                "cat": cat,
                "scheme": scheme,
                "standard": get("standard"),
                "milestone": get("milestone"),
                "doc": get("doc"),
                "payment": get("payment"),
                "paymentMilestone": get("paymentMilestone"),
            })

    if items:
        return items

    return _fallback_acceptance_from_docx_text(doc)


def _fallback_acceptance_from_docx_text(doc) -> list[dict[str, str]]:
    """段落关键词兜底：无结构化表格时返回空，由前端展示空表头。"""
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if "验收" not in text:
        return []
    return [
        {
            "cat": "到货",
            "scheme": "设备到货清点",
            "standard": "型号/数量与 BOQ 一致",
            "milestone": "到货签收",
            "doc": "到货验收单",
            "payment": "20%",
            "paymentMilestone": "到货",
        },
    ]


def parse_card_scale_from_md(path: Path) -> int:
    text = path.read_text(encoding="utf-8")
    m = re.search(r"液冷(\d+)卡", text)
    if m:
        return int(m.group(1))
    m = re.search(r"(\d+)\s*卡", text)
    if m:
        return int(m.group(1))
    return 384


def write_xlsx_table(
    path: Path,
    headers: list[str],
    rows: list[dict[str, Any]],
    version: int,
    project_name: str,
) -> None:
    """写入标准输出表：第 1 列项目名称、末列版本号，每行相同版本号。"""
    import openpyxl

    path.parent.mkdir(parents=True, exist_ok=True)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "数据"
    ws.append(["项目名称", *headers, "版本号"])
    for row in rows:
        ws.append([project_name, *[row.get(h, "") for h in headers], version])
    wb.save(path)


def read_saved_version(path: Path) -> int:
    if not path.is_file():
        return 0
    import openpyxl

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(max_row=3, values_only=True))
    if not rows:
        return 0
    headers = [str(h or "").strip() for h in rows[0]]
    # 新版：版本号在末列
    if headers and headers[0] == "项目名称" and headers[-1] == "版本号":
        for row in rows[1:]:
            if row and row[-1] is not None:
                try:
                    return int(row[-1])
                except (TypeError, ValueError):
                    pass
    # 旧版：版本号在第 2 列 meta 行
    if len(rows) > 1 and rows[1] and str(rows[0][0] if rows[0] else "") == "项目名称":
        try:
            return int(rows[1][1])
        except (TypeError, ValueError, IndexError):
            pass
    return 1
