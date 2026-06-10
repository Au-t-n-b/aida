"""
智慧工勘 v4 — 报告模板填充

职责: 填充工勘报告.docx 模板（9个表格）
错误码前缀: SS-RB

注意: 当前实现基于 python-docx 操作。
照片功能在有实际模板和照片数据时才能完整测试。
"""
from __future__ import annotations

import os
from typing import Optional

from .logger import log_info, log_warn
from .types import (
    AssessmentResult,
    AssessmentValue,
    IssueItem,
    LLMCallable,
    ProjectMeta,
    RiskItem,
    SurveyResultRow,
)


# ──────────────────────────────────────────────
# 异常定义
# ──────────────────────────────────────────────

class ReportBuildError(Exception):
    def __init__(self, code: str, message: str):
        self.code = code
        self.message = message
        super().__init__(f"[{code}] {message}")


# ──────────────────────────────────────────────
# 评估值符号映射
# ──────────────────────────────────────────────

ASSESSMENT_SYMBOLS = {
    "满足": "✅",
    "不满足": "❌",
    "不涉及": "◯",
    "未勘测": "🟡",
    "无法识别": "❓",
}


# ──────────────────────────────────────────────
# 公开接口
# ──────────────────────────────────────────────

def build_report(
    template_path: str,
    output_dir: str,
    project_meta: ProjectMeta,
    survey_results: list[SurveyResultRow],
    assessment_results: list[AssessmentResult],
    issue_list: list[IssueItem],
    risk_list: list[RiskItem],
    photo_map: Optional[dict[int, list[bytes]]] = None,
    llm_call: Optional[LLMCallable] = None,
) -> str:
    """
    填充报告模板，生成工勘报告。

    返回:
        生成的报告文件路径
    """
    try:
        from docx import Document
    except ImportError:
        raise ReportBuildError("SS-RB-E-001", "python-docx 未安装，无法生成报告")

    if not os.path.exists(template_path):
        raise ReportBuildError("SS-RB-E-001", f"报告模板文件不存在: {template_path}")

    doc = Document(template_path)

    # 验证表格数量
    if len(doc.tables) < 9:
        raise ReportBuildError(
            "SS-RB-E-002",
            f"报告模板表格数量不足: 期望≥9, 实际={len(doc.tables)}",
        )

    # 依次填充各表格
    _fill_table0_basic_info(doc, project_meta)
    _fill_table1_photos(doc, survey_results, photo_map or {}, llm_call)
    _fill_table2_overview(doc, survey_results, assessment_results)
    _fill_table3_4_detail(doc, survey_results, assessment_results)
    _fill_table5_risks(doc, risk_list)
    _fill_table6_issues(doc, issue_list)
    _fill_table7_unsurveyed(doc, survey_results, assessment_results)
    _fill_table8_unrecognized(doc, survey_results, assessment_results)

    # 替换段落占位符
    stats = _compute_statistics(assessment_results)
    replacements = {
        "{项目名称}": project_meta.project_name,
        "{机房名称}": project_meta.room_name,
        "{来自总计检查项数}": str(len(survey_results)),
        "{满足项数量}": str(stats["满足"]),
        "{不满足项}": str(stats["不满足"]),
        "{不涉及项}": str(stats["不涉及"]),
        "{未勘测项数量}": str(stats["未勘测"]),
        "{无法识别项数量}": str(stats["无法识别"]),
    }
    _replace_paragraph_placeholders(doc, replacements)

    # 保存
    os.makedirs(output_dir, exist_ok=True)
    filename = (
        f"{project_meta.activity_id}_{project_meta.project_name}_"
        f"{project_meta.room_name}_工勘报告.docx"
    )
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)

    log_info("report_builder", "build_report", f"报告已生成: {output_path}")
    return output_path


def get_report_statistics(
    survey_results: list[SurveyResultRow],
    assessment_results: list[AssessmentResult],
) -> dict:
    """获取报告摘要统计信息。"""
    stats = _compute_statistics(assessment_results)
    return {
        "total": len(survey_results),
        "statistics": stats,
        "scenes": list(set(r.get("细分场景", "") for r in survey_results if r.get("细分场景"))),
    }


# ──────────────────────────────────────────────
# 内部函数 — 表格填充
# ──────────────────────────────────────────────

def _fill_table0_basic_info(doc, project_meta: ProjectMeta) -> None:
    """Table[0]: 基本信息表 — 替换占位符。"""
    table = doc.tables[0]
    replacements = {
        "{项目名称}": project_meta.project_name,
        "{勘测日期}": project_meta.survey_date,
        "{勘测房间}": project_meta.room_name,
        "{勘测人员}": project_meta.surveyor,
    }
    _replace_table_placeholders(table, replacements)


def _fill_table1_photos(
    doc,
    survey_results: list[SurveyResultRow],
    photo_map: dict[int, list[bytes]],
    llm_call: Optional[LLMCallable],
) -> None:
    """
    Table[1]: 8×2 照片网格。
    奇数行: 照片说明（LLM 生成或默认使用勘测要素名）
    偶数行: 插入照片
    最多填充8张照片。
    """
    if not photo_map:
        return

    if len(doc.tables) < 2:
        return

    table = doc.tables[1]

    # 收集有照片的行，优先不满足项
    photo_candidates = []
    for row_idx, photos in sorted(photo_map.items()):
        if row_idx < len(survey_results):
            result = survey_results[row_idx]
            ai_val = result.get("AI评估结果", "")
            priority = 0 if "不满足" in ai_val else (1 if "满足" in ai_val else 2)
            for photo_bytes in photos:
                photo_candidates.append((priority, row_idx, photo_bytes, result))

    photo_candidates.sort(key=lambda x: x[0])
    selected = photo_candidates[:8]

    if not selected:
        return

    import io
    try:
        from docx.shared import Cm
    except ImportError:
        return

    for i, (_, row_idx, photo_bytes, result) in enumerate(selected):
        desc_row_idx = i * 2
        photo_row_idx = i * 2 + 1

        if desc_row_idx >= len(table.rows) or photo_row_idx >= len(table.rows):
            break

        # 生成照片说明
        factor = result.get("勘测要素", "")
        item = result.get("项目", "")
        description = f"{factor}-{item}" if factor and item else (factor or item or f"照片{i+1}")

        if llm_call:
            try:
                prompt = (
                    f"请为以下工勘照片生成≤15字的简洁描述：\n"
                    f"勘测要素: {factor}\n项目: {item}\n"
                    f"检查内容: {result.get('检查内容', '')}"
                )
                llm_desc = llm_call(prompt)
                if llm_desc and len(llm_desc.strip()) <= 20:
                    description = llm_desc.strip()
            except Exception:
                pass

        # 填充说明行
        desc_cell = table.rows[desc_row_idx].cells[0] if table.rows[desc_row_idx].cells else None
        if desc_cell:
            desc_cell.text = description

        # 填充照片行
        photo_cell = table.rows[photo_row_idx].cells[0] if table.rows[photo_row_idx].cells else None
        if photo_cell:
            try:
                paragraph = photo_cell.paragraphs[0] if photo_cell.paragraphs else photo_cell.add_paragraph()
                run = paragraph.add_run()
                image_stream = io.BytesIO(photo_bytes)
                run.add_picture(image_stream, width=Cm(6))
            except Exception:
                photo_cell.text = f"[照片{i+1}]"


def _fill_table2_overview(
    doc,
    survey_results: list[SurveyResultRow],
    assessment_results: list[AssessmentResult],
) -> None:
    """Table[2]: 概览统计表。"""
    table = doc.tables[2]

    # 按 (细分场景, 勘测要素) 分组统计
    groups = _group_by_scene_factor(survey_results, assessment_results)

    # 清空模板数据行（保留表头）
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    # 写入数据
    for group in groups:
        row = table.add_row()
        cells = row.cells
        cells[0].text = group["scene"]
        cells[1].text = group["factor"]
        cells[2].text = str(group["total"])
        cells[3].text = str(group["满足"])
        cells[4].text = str(group["不满足"])
        cells[5].text = str(group["不涉及"])
        cells[6].text = str(group["未勘测"])
        cells[7].text = str(group["无法识别"])

    # 添加总计行
    totals = _compute_statistics(assessment_results)
    total_row = table.add_row()
    cells = total_row.cells
    cells[0].text = "总计"
    cells[1].text = ""
    cells[2].text = str(len(survey_results))
    cells[3].text = str(totals["满足"])
    cells[4].text = str(totals["不满足"])
    cells[5].text = str(totals["不涉及"])
    cells[6].text = str(totals["未勘测"])
    cells[7].text = str(totals["无法识别"])


def _fill_table3_4_detail(
    doc,
    survey_results: list[SurveyResultRow],
    assessment_results: list[AssessmentResult],
) -> None:
    """Table[3]~[4+]: 分场景详细评估表。"""
    if len(doc.tables) < 4:
        return

    table = doc.tables[3]

    # 清空模板行
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    # 按细分场景分组
    scene_groups = {}
    for i, (result, ar) in enumerate(zip(survey_results, assessment_results)):
        scene = result.get("细分场景", "其他")
        if scene not in scene_groups:
            scene_groups[scene] = []
        scene_groups[scene].append((result, ar))

    seq = 0
    for scene, items in scene_groups.items():
        for result, ar in items:
            seq += 1
            row = table.add_row()
            cells = row.cells
            cells[0].text = str(seq)
            cells[1].text = result.get("勘测要素", "")
            cells[2].text = result.get("项目", "")
            cells[3].text = result.get("检查内容", "")
            cells[4].text = result.get("最新检查结果", "")
            cells[5].text = ASSESSMENT_SYMBOLS.get(ar.conclusion.value, "")
            cells[6].text = ar.defect_description if ar.conclusion in (
                AssessmentValue.UNSATISFIED, AssessmentValue.UNRECOGNIZABLE
            ) else ""


def _fill_table5_risks(doc, risk_list: list[RiskItem]) -> None:
    """Table[5]: 风险表。"""
    if len(doc.tables) < 6:
        return

    table = doc.tables[5]

    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    for i, risk in enumerate(risk_list, 1):
        row = table.add_row()
        cells = row.cells
        cells[0].text = str(i)
        cells[1].text = risk.get("风险类型描述", "")
        cells[2].text = risk.get("风险细项描述", "")
        cells[3].text = risk.get("风险影响", "")
        cells[4].text = risk.get("建议措施", "")
        cells[5].text = risk.get("备注", "")


def _fill_table6_issues(doc, issue_list: list[IssueItem]) -> None:
    """Table[6]: 问题整改待办。"""
    if len(doc.tables) < 7:
        return

    table = doc.tables[6]

    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    for i, issue in enumerate(issue_list, 1):
        row = table.add_row()
        cells = row.cells
        cells[0].text = str(i)
        cells[1].text = issue.get("问题描述", "")
        cells[2].text = issue.get("整改建议", "")
        cells[3].text = issue.get("责任人", "")
        cells[4].text = issue.get("计划关闭时间", "")
        cells[5].text = issue.get("备注", "")


def _fill_table7_unsurveyed(
    doc,
    survey_results: list[SurveyResultRow],
    assessment_results: list[AssessmentResult],
) -> None:
    """Table[7]: 遗留未勘测待办。"""
    if len(doc.tables) < 8:
        return

    table = doc.tables[7]

    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    seq = 0
    for result, ar in zip(survey_results, assessment_results):
        if ar.conclusion == AssessmentValue.NOT_SURVEYED:
            seq += 1
            row = table.add_row()
            cells = row.cells
            cells[0].text = str(seq)
            cells[1].text = f"{result.get('项目', '')} - {result.get('检查内容', '')}"
            cells[2].text = ""
            cells[3].text = ""
            cells[4].text = ""


def _fill_table8_unrecognized(
    doc,
    survey_results: list[SurveyResultRow],
    assessment_results: list[AssessmentResult],
) -> None:
    """Table[8]: 遗留未评估待办。"""
    if len(doc.tables) < 9:
        return

    table = doc.tables[8]

    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    seq = 0
    for result, ar in zip(survey_results, assessment_results):
        if ar.conclusion == AssessmentValue.UNRECOGNIZABLE:
            seq += 1
            row = table.add_row()
            cells = row.cells
            cells[0].text = str(seq)
            cells[1].text = f"{result.get('项目', '')} - {result.get('检查内容', '')}"
            cells[2].text = ""
            cells[3].text = ""
            cells[4].text = ar.defect_description


# ──────────────────────────────────────────────
# 内部工具
# ──────────────────────────────────────────────

def _compute_statistics(assessment_results: list[AssessmentResult]) -> dict[str, int]:
    """计算五值统计。"""
    stats = {v.value: 0 for v in AssessmentValue}
    for ar in assessment_results:
        stats[ar.conclusion.value] += 1
    return stats


def _group_by_scene_factor(
    survey_results: list[SurveyResultRow],
    assessment_results: list[AssessmentResult],
) -> list[dict]:
    """按 (细分场景, 勘测要素) 分组统计。"""
    groups: dict[tuple[str, str], dict] = {}

    for result, ar in zip(survey_results, assessment_results):
        scene = result.get("细分场景", "其他")
        factor = result.get("勘测要素", "其他")
        key = (scene, factor)

        if key not in groups:
            groups[key] = {
                "scene": scene,
                "factor": factor,
                "total": 0,
                "满足": 0,
                "不满足": 0,
                "不涉及": 0,
                "未勘测": 0,
                "无法识别": 0,
            }

        groups[key]["total"] += 1
        groups[key][ar.conclusion.value] += 1

    return list(groups.values())


def _replace_table_placeholders(table, replacements: dict[str, str]) -> None:
    """替换表格中的占位符。"""
    for row in table.rows:
        for cell in row.cells:
            for key, val in replacements.items():
                if key in cell.text:
                    cell.text = cell.text.replace(key, val)


def _replace_paragraph_placeholders(doc, replacements: dict[str, str]) -> None:
    """遍历文档所有段落替换占位符。"""
    for paragraph in doc.paragraphs:
        for key, val in replacements.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, val)
