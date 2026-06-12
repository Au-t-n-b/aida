#!/usr/bin/env python3
"""
测试用例后处理引擎
==================
从测试用例 .docx 或 .md 全文提取结构化 JSON。
分类规则来自 SKILL.md §1.3。

用法:
    import testcase_post
    result = testcase_post.extract("测试用例-全文解析.md")
    result = testcase_post.extract("测试用例.docx")    # 直解析
"""

import json
import re
import os
from datetime import datetime


# ============================================================
# 测试用例分类规则 — SKILL.md §1.3
# ============================================================

CATEGORY_RULES = {
    "硬件设备": {
        "keywords": ["设备到货", "设备验收", "Pod", "物理设备", "硬件", "上架", "配电",
                      "指示灯", "BIOS", "RAID", "冗余", "告警上报", "机柜",
                      "开箱", "外包装", "电源框", "风扇", "硬盘", "滑轨"],
        "exclude": ["软件", "License", "iBMC", "OS", "操作系统"],
    },
    "软件部分": {
        "keywords": ["License", "激活", "商用永久授权", "OS", "操作系统", "iBMC",
                      "驱动", "固件", "NPU驱动", "ZTP", "软件版本", "一键式收集日志",
                      "管理模块", "管理页面", "Web界面", "远程控制"],
    },
    "服务部分": {
        "keywords": ["网络", "灵衢", "BGP", "光模块", "光口", "lane", "CPLD",
                      "电子标签", "温度管理", "CloudOps", "集群", "集合通信",
                      "模型训练", "PyTorch", "摸高", "稳定性", "PRBS", "压力",
                      "HCCS", "RoCE", "参数面", "健康检查", "健康状态",
                      "带宽性能", "读写性能", "算力性能", "连通性", "互通性"],
    }
}


def classify(text: str) -> str:
    """关键词投票分类。"""
    scores = {}
    for cat, rules in CATEGORY_RULES.items():
        score = sum(1 for kw in rules["keywords"] if kw.lower() in text.lower())
        score -= sum(2 for ekw in rules.get("exclude", []) if ekw.lower() in text.lower())
        scores[cat] = score
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "服务部分"


def classify_layer(text: str) -> str:
    """服务部分三层分层: 下路/中路/上路。"""
    lower = sum(1 for kw in ["网络", "灵衢", "BGP", "光模块", "光口", "lane", "CPLD",
                               "电子标签", "温度管理", "PRBS", "压力", "互通性"]
                if kw.lower() in text.lower())
    middle = sum(1 for kw in ["CloudOps", "集群", "集合通信", "健康检查", "健康状态",
                                "带宽性能", "读写性能", "算力性能", "连通性", "HCCS"]
                 if kw.lower() in text.lower())
    upper = sum(1 for kw in ["模型训练", "PyTorch", "摸高", "稳定性测试", "qwen"]
                if kw.lower() in text.lower())
    if upper >= max(lower, middle): return "上路-业务应用"
    if middle >= max(lower, upper): return "中路-平台建立"
    return "下路-计算存储网络"


# ============================================================
# 主入口
# ============================================================

def extract(path: str) -> dict:
    """从 .docx 或 .md 提取测试用例结构化数据。"""
    if path.endswith(".docx"):
        return _from_docx(path)
    return _from_md(path)


# ============================================================
# Docx 直解析
# ============================================================

def _from_docx(docx_path: str) -> dict:
    from docx import Document

    doc = Document(docx_path)
    source = {"filename": os.path.basename(docx_path), "total_cases": 0,
              "parsed_at": datetime.now().isoformat()}

    # TOC → case_id 映射
    num_map = {}
    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        if style.startswith("toc "):
            m = re.match(r'([\d.]+)\s+(.+?)\t', para.text.strip())
            if m: num_map[m.group(2).strip()] = m.group(1)

    cases = []
    cur = None
    content = []
    para_idx = table_idx = 0
    ch = sec = cnum = 0
    cur_sec = ""
    in_rules = False
    ctx = None  # preconditions | steps | expected | passfail

    for child in doc.element.body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            if para_idx >= len(doc.paragraphs): para_idx += 1; continue
            para = doc.paragraphs[para_idx]; para_idx += 1
            text = para.text.strip()
            style = para.style.name if para.style else ""

            if style.startswith("toc "): continue
            if not text: continue

            if style.startswith("Heading 1"): ch += 1; sec = 0; continue
            if style.startswith("Heading 2"): sec += 1; cnum = 0; cur_sec = text; continue

            if style.startswith("Heading 3"):
                if cur: _finish(cur, content); cases.append(cur)
                cnum += 1
                cid = num_map.get(text, f"{ch}.{sec}.{cnum}")
                cur = _new_case(cid, f"§{ch}.{sec} {cur_sec}", text)
                content = [f"# {text}"]
                in_rules = False; ctx = None
                continue

            if cur:
                content.append(text)
                if "规则与逻辑" in text: in_rules = True; ctx = None
                elif "用例摘要" in text: in_rules = False
                elif in_rules:
                    if "前置条件" in text:
                        ctx = "preconditions"
                        v = re.sub(r'^.*?前置条件[（(].*?[）)]', '', text).strip()
                        if v: cur["preconditions"].append(v)
                    elif "测试步骤" in text: ctx = "steps"
                    elif "预期结果" in text: ctx = "expected"
                    elif "通过/失败判定" in text: ctx = "passfail"
                    elif ctx == "preconditions": cur["preconditions"].append(text)
                    elif ctx == "steps":
                        m = re.match(r'(\d+)\.\s*(.+)', text)
                        if m: cur["test_steps"].append(m.group(2).strip())
                    elif ctx == "expected":
                        m = re.match(r'(\d+)\.\s*(.+)', text)
                        if m: cur["expected_result"].append(m.group(2).strip())
                        else: cur["expected_result"].append(text)
                    elif ctx == "passfail":
                        pm = re.search(r'通过[：:]\s*(.+)', text)
                        if pm: cur["pass_fail_criteria"]["pass"] = pm.group(1).strip()
                        fm = re.search(r'失败[：:]\s*(.+)', text)
                        if fm: cur["pass_fail_criteria"]["fail"] = fm.group(1).strip()

        elif tag == 'tbl':
            if table_idx >= len(doc.tables): table_idx += 1; continue
            table = doc.tables[table_idx]; table_idx += 1
            rows = [[c.text.strip().replace('\n', ' ') for c in row.cells] for row in table.rows]
            if not rows or not rows[0]: continue

            hdr = rows[0][0] if rows[0] else ""
            if cur:
                content.append("")
                for row in rows: content.append("| " + " | ".join(row) + " |")
                content.append("")

            if "实体名" in hdr: cur["entities"] = _parse_entities(rows)
            elif "源实体" in hdr: cur["relations"] = _parse_relations(rows)
            else: _parse_summary(rows, cur)

    if cur: _finish(cur, content); cases.append(cur)
    source["total_cases"] = len(cases)
    return {"source": source, "cases": cases}


# ============================================================
# Markdown 回退解析
# ============================================================

def _from_md(md_path: str) -> dict:
    with open(md_path, "r", encoding="utf-8") as f:
        text = f.read()
    text = re.sub(r'\s*\[§[\d.]*(?:-p\d+)?\]', '', text)

    blocks = re.split(r'\n(?=# 用例 )', text)
    source = {"filename": os.path.basename(md_path), "total_cases": 0,
              "parsed_at": datetime.now().isoformat()}
    cases = []

    for block in blocks:
        m = re.match(r'# 用例\s+(.+)', block)
        if not m: continue
        cid = re.sub(r'\s*\[.*?\]\s*$', '', m.group(1).strip())

        case = {
            "case_id": cid,
            "chapter": _derive_ch(cid),
            "test_purpose": _field(block, "测试目的"),
            "preconditions": _section(block, "前置条件"),
            "test_steps": _numbered(block, "测试步骤"),
            "expected_result": _numbered(block, "预期结果"),
            "pass_fail_criteria": _passfail(block),
            "entities": _etable(block),
            "relations": _rtable(block),
            "test_network": _field(block, "测试组网") or "NA",
            "remarks": _field(block, "备注") or "",
            "category": classify(block),
            "layer": classify_layer(block) if classify(block) == "服务部分" else None,
            "content": block.strip()
        }
        cases.append(case)

    source["total_cases"] = len(cases)
    return {"source": source, "cases": cases}


# ============================================================
# 内部工具
# ============================================================

def _new_case(cid, ch, purpose):
    return {"case_id": cid, "chapter": ch, "test_purpose": purpose,
            "preconditions": [], "test_steps": [], "expected_result": [],
            "pass_fail_criteria": {"pass": "", "fail": ""},
            "entities": [], "relations": [],
            "test_network": "NA", "remarks": "",
            "category": "", "layer": None, "content": ""}

def _finish(case, lines):
    case["content"] = "\n".join(lines)
    case["category"] = classify(case["content"])
    if case["category"] == "服务部分": case["layer"] = classify_layer(case["content"])

def _derive_ch(cid):
    parts = cid.split(".")
    m = {
        "1.1":"§1.1 硬件外观检查","1.2":"§1.2 iBMC管理功能","1.3":"§1.3 健康状态检查",
        "1.4":"§1.4 冗余与告警","1.5":"§1.5 压力测试","1.6":"§1.6 软件版本检查",
        "1.7":"§1.7 模型训练","2.1":"§2.1 PRBS压测","3.1":"§3.1 超节点健康检查",
        "3.2":"§3.2 CloudOps集合通信","3.3":"§3.3 CloudOps基础性能",
        "3.4":"§3.4 模型训练摸高","3.5":"§3.5 模型稳定性"
    }
    return m.get(f"{parts[0]}.{parts[1]}", f"§{cid}") if len(parts)>=2 else f"§{cid}"

def _field(block, name):
    m = re.search(rf'{name}\s*\|\s*(.+?)(?:\s*\||\s*$)', block)
    return re.sub(r'\s*\|\s*$', '', m.group(1).strip()) if m else ""

def _section(block, name):
    m = re.search(rf'\*\*{name}\*\*(?:[（(].*?[）)])?\s*\n\s*((?:.(?!\n\s*-\s*\*\*))+.)', block, re.DOTALL)
    return [l.strip() for l in m.group(1).strip().split('\n') if l.strip()] if m else []

def _numbered(block, name):
    m = re.search(rf'\*\*{name}\*\*(?:[（(].*?[）)])?\s*\n\s*((?:.(?!\n\s*-\s*\*\*))+.)', block, re.DOTALL)
    if not m: return []
    items = []
    for line in m.group(1).split('\n'):
        sm = re.match(r'(\d+)\.\s*(.+)', line.strip())
        if sm: items.append(re.sub(r'\s+$', '', sm.group(2)))
        elif line.strip(): items.append(line.strip())
    return items

def _passfail(block):
    pf = {"pass":"","fail":""}
    m = re.search(r'\*\*通过/失败判定\*\*\s*\n\s*((?:.(?!\n\s*##))+.)', block, re.DOTALL)
    if m:
        pm = re.search(r'通过\*\*[：:]\s*(.+?)(?=\n\s*-\s*\*\*失败|\Z)', m.group(1), re.DOTALL)
        if pm: pf["pass"] = pm.group(1).strip()
        fm = re.search(r'失败\*\*[：:]\s*(.+)', m.group(1), re.DOTALL)
        if fm: pf["fail"] = fm.group(1).strip()
    return pf

def _etable(block):
    rows = []
    m = re.search(r'实体[（(].*?[）)]\s*\n\|.*?\n\|[-| ]+\n((?:\|.+\n)+)', block)
    if m:
        for line in m.group(1).strip().split('\n'):
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if len(cells) >= 4:
                rows.append({"name":cells[0],"type":cells[1],
                             "attributes":[a.strip() for a in cells[2].split('、')] if cells[2] else [],
                             "description":cells[3] if len(cells)>3 else ""})
    return rows

def _rtable(block):
    rows = []
    m = re.search(r'关系[（(].*?[）)]\s*\n\|.*?\n\|[-| ]+\n((?:\|.+\n)+)', block)
    if m:
        for line in m.group(1).strip().split('\n'):
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if len(cells) >= 4:
                rows.append({"source":cells[0],"relation":cells[1],"target":cells[2],"cardinality":cells[3]})
    return rows

def _parse_entities(rows):
    ents = []
    for cells in rows:
        if len(cells) >= 4 and cells[0] not in ("实体名","名称",""):
            ents.append({"name":cells[0],"type":cells[1] if len(cells)>1 else "",
                         "attributes":[a.strip() for a in cells[2].split('、')] if len(cells)>2 and cells[2] else [],
                         "description":cells[3] if len(cells)>3 else ""})
    return ents

def _parse_relations(rows):
    rels = []
    for cells in rows:
        if len(cells) >= 4 and cells[0] not in ("源实体","实体",""):
            rels.append({"source":cells[0],"relation":cells[1] if len(cells)>1 else "",
                         "target":cells[2] if len(cells)>2 else "","cardinality":cells[3] if len(cells)>3 else ""})
    return rels

def _parse_summary(rows, case):
    if not case: return
    for cells in rows:
        if len(cells) < 2: continue
        k = cells[0].strip().replace(" ", "")
        v = cells[1].strip()
        if not v: continue
        if k in ("用例编号","编号"): case["case_id"] = v
        elif k in ("测试目的","目的"): case["test_purpose"] = v
        elif k in ("测试组网","组网"): case["test_network"] = v
        elif "备注" in k or k == "注": case["remarks"] = v
        elif "前置条件" in k: case["preconditions"] = _split(v)
        elif "测试步骤" in k: case["test_steps"] = _split(v)
        elif "预期结果" in k: case["expected_result"] = _split(v)
        elif "通过" in k and "失败" not in k: case["pass_fail_criteria"]["pass"] = v
        elif "失败" in k: case["pass_fail_criteria"]["fail"] = v

def _split(text):
    if not text.strip(): return []
    parts = re.split(r'(?:^|\s+)(\d+)\.\s*', text)
    if len(parts) > 2:
        return [parts[i+1].strip().rstrip('。') for i in range(1, len(parts), 2) if i+1 < len(parts)]
    parts = [p.strip().rstrip('。') for p in re.split(r'[。；]\s+', text) if p.strip()]
    return parts if len(parts) > 1 else [text.strip()]


# ============================================================
# CLI
# ============================================================

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("用法: python testcase-post.py <file.docx|file.md>")
        sys.exit(1)
    result = extract(sys.argv[1])
    print(json.dumps(result, ensure_ascii=False, indent=2))
