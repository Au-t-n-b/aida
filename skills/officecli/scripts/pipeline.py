#!/usr/bin/env python3
"""
DSXW 验收解析流水线 (Orchestrator)
===================================
三层架构:
  1. 全文提取: Docx → Markdown (python-docx / pandoc)
  2. 规则匹配: scripts.py 基于 SKILL.md 规则从全文提取结构化 JSON
  3. 输出文件: {name}-全文解析.md + {name}-结构化输出.json

用法:
    python pipeline.py                          # 跑全部
    python pipeline.py --pipeline proposal      # 仅服务建议书
    python pipeline.py --pipeline testcases     # 仅测试用例
"""

import json
import os
import re
import sys
import shutil
from pathlib import Path
from datetime import datetime

BASE_DIR = Path(__file__).resolve().parent

# 导入提取引擎 (需先加入 path)
sys.path.insert(0, str(BASE_DIR))
import testcase_post
import service_advice_post

# ============================================================
# Docx → Markdown 全文转换
# ============================================================

def docx_to_markdown(docx_path: str) -> str:
    """将 Docx 完整转换为 Markdown 全文 (python-docx 方案)。

    优先使用 pandoc (质量更高), 不可用时回退 python-docx。
    """
    import subprocess

    # 尝试 pandoc
    try:
        result = subprocess.run(
            ["pandoc", docx_path, "-t", "markdown", "--wrap=none"],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode == 0 and len(result.stdout) > 1000:
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # 回退 python-docx
    try:
        from docx import Document
        doc = Document(docx_path)
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                lines.append("")
                continue
            style = para.style.name if para.style else ""
            if style.startswith("Heading 1"):
                lines.append(f"# {text}")
            elif style.startswith("Heading 2"):
                lines.append(f"## {text}")
            elif style.startswith("Heading 3"):
                lines.append(f"### {text}")
            elif style.startswith("Heading"):
                lines.append(f"#### {text}")
            else:
                lines.append(text)

        for table in doc.tables:
            lines.append("")
            if table.rows:
                lines.append("| " + " | ".join(
                    cell.text.strip().replace("\n", " ") for cell in table.rows[0].cells
                ) + " |")
                lines.append("|" + "|".join(["---"] * len(table.rows[0].cells)) + "|")
                for row in table.rows[1:]:
                    cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    lines.append("| " + " | ".join(cells) + " |")
            lines.append("")
        return "\n".join(lines)
    except ImportError:
        raise RuntimeError("需要 pandoc 或 python-docx: pip install python-docx")


def add_paragraph_numbers(md_text: str) -> str:
    """为 Markdown 全文添加层级段落序号。

    序号反映文档章节结构，追加在行尾:
      # 标题        → [§1]
      ## 子标题      → [§1.1]
      ### 子子标题   → [§1.1.1]
      正文段落       → [§1.1-p1], [§1.1-p2] ...
      表格行         → 继承当前节序号

    空白行保留不编号。
    """
    lines = md_text.split('\n')
    result = []
    # 层级计数器: h1, h2, h3, h4 ...
    stack = [0, 0, 0, 0, 0, 0]
    para_n = 0  # 当前节内段落计数

    for line in lines:
        stripped = line.strip()
        if not stripped:
            result.append(line)
            continue

        # 检测标题层级
        heading_match = re.match(r'^(#{1,6})\s+(.+)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            # 重置当前层及以下
            stack[level - 1] += 1
            for i in range(level, len(stack)):
                stack[i] = 0
            para_n = 0  # 新节重置段落计数

            # 构建节号: §1 / §1.1 / §1.1.1
            parts = [str(stack[i]) for i in range(level) if stack[i] > 0]
            sec_id = ".".join(parts)
            result.append(f"{line} [§{sec_id}]")
        else:
            para_n += 1
            # 当前节号
            parts = [str(stack[i]) for i in range(len(stack)) if stack[i] > 0]
            sec_id = ".".join(parts) if parts else "0"
            result.append(f"{line} [§{sec_id}-p{para_n}]")

    return '\n'.join(result)


# ============================================================
# Pipeline 1: 服务建议书
# ============================================================

def run_proposal_pipeline(proposal_docx: str, out_dir: Path):
    """服务建议书 → 全文解析.md + 验收表.json"""
    print(f"\n{'='*60}")
    print("Pipeline 1: 服务建议书 → 全文解析 + 验收表")
    print(f"{'='*60}")

    # Step 0: Docx → Markdown 全文解析
    print(f"\n  [Step 0] Docx → Markdown 全文解析")
    print(f"    Input:  {os.path.basename(proposal_docx)}")
    md_text = docx_to_markdown(proposal_docx)
    md_text = add_paragraph_numbers(md_text)  # 添加段落序号
    fulltext_path = out_dir / "服务建议书-全文解析.md"
    with open(fulltext_path, "w", encoding="utf-8") as f:
        f.write(md_text)
    print(f"    [OK] {fulltext_path.name} ({len(md_text)/1024:.0f} KB)")

    # Step 1: 从全文提取验收表 (scripts.py)
    print(f"\n  [Step 1] service-advice-post.py 提取验收表")
    table = service_advice_post.extract(str(fulltext_path))
    for item in table:
        print(f"    [{item['id']}] {item['category']}")

    # Step 2: 写结构化输出
    print(f"\n  [Step 2] 写结构化输出")
    json_path = out_dir / "服务验收表-结构化输出.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(table, f, ensure_ascii=False, indent=2)
    print(f"    [OK] {json_path.name}")


# ============================================================
# Pipeline 2: 测试用例
# ============================================================

def run_testcases_pipeline(testcase_src: str, out_dir: Path):
    """测试用例 → 全文解析.md + 结构化.json"""
    print(f"\n{'='*60}")
    print("Pipeline 2: 测试用例 → 全文解析 + 结构化提取")
    print(f"{'='*60}")

    is_docx = testcase_src.endswith(".docx")
    fulltext_path = out_dir / "测试用例-全文解析.md"

    # Step 0: 全文解析
    print(f"\n  [Step 0] 全文解析")
    print(f"    Input:  {os.path.basename(testcase_src)} ({'.docx' if is_docx else '.md'})")

    if is_docx:
        md_text = docx_to_markdown(testcase_src)
    else:
        with open(testcase_src, "r", encoding="utf-8") as f:
            md_text = f.read()

    md_text = add_paragraph_numbers(md_text)
    with open(fulltext_path, "w", encoding="utf-8") as f:
        f.write(md_text)
    print(f"    [OK] {fulltext_path.name} ({len(md_text)/1024:.0f} KB)")

    # Step 1: scripts.py 提取结构化数据
    # 如果是 .docx，优先直接用 .docx 解析（质量更高）
    extract_input = testcase_src if is_docx else str(fulltext_path)
    print(f"\n  [Step 1] testcase-post.py 提取结构化数据")
    print(f"    Extract from: {os.path.basename(extract_input)}")
    result = testcase_post.extract(extract_input)
    cases = result["cases"]
    print(f"    Parsed: {len(cases)} cases")

    cats = {}
    for c in cases:
        cat = c.get("category", "unknown")
        cats[cat] = cats.get(cat, 0) + 1
    for cat, count in cats.items():
        print(f"      {cat}: {count}")

    # Step 2: 写结构化输出
    print(f"\n  [Step 2] 写结构化输出")
    json_path = out_dir / "测试用例-结构化输出.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    print(f"    [OK] {json_path.name} ({os.path.getsize(json_path)/1024:.0f} KB)")


# ============================================================
# Main
# ============================================================

def main():
    import argparse
    parser = argparse.ArgumentParser(description="DSXW 验收解析流水线")
    parser.add_argument("--pipeline", choices=["proposal", "testcases", "all"], default="all",
                        help="执行哪条流水线")
    parser.add_argument("--proposal-docx", default=None,
                        help="服务建议书 Docx 路径")
    parser.add_argument("--testcase-src", default=None,
                        help="测试用例源文件 (.docx 或 .md)")
    parser.add_argument("--out-dir", default=None, help="输出目录")
    args = parser.parse_args()

    out_dir = Path(args.out_dir) if args.out_dir else BASE_DIR
    run_all = args.pipeline == "all"

    # --- Pipeline 1: 服务建议书 ---
    if run_all or args.pipeline == "proposal":
        proposal = args.proposal_docx or str(
            BASE_DIR / "DSXW算力服务器项目_技术建议书(含算力服务器、网络、存储、机房租赁) -不含人员信息.docx"
        )
        if os.path.exists(proposal):
            run_proposal_pipeline(proposal, out_dir)
        else:
            print(f"[WARN] 服务建议书不存在: {proposal}")

    # --- Pipeline 2: 测试用例 ---
    if run_all or args.pipeline == "testcases":
        # 优先用 .docx (scripts.py 有直接解析能力)
        existing_docx = BASE_DIR / "A3超节点验收测试用例-液冷场景V1.2.docx"
        existing_md = BASE_DIR / "A3超节点验收测试用例.md"

        testcase_src = args.testcase_src
        if not testcase_src:
            testcase_src = str(existing_docx) if existing_docx.exists() else str(existing_md)

        if os.path.exists(testcase_src):
            run_testcases_pipeline(testcase_src, out_dir)
        else:
            print(f"[WARN] 测试用例源文件不存在")

    # --- 汇总 ---
    print(f"\n{'='*60}")
    print("流水线完成!")
    print(f"{'='*60}")
    print(f"\n输出目录: {out_dir}")
    for f in sorted(out_dir.glob("*")):
        name = f.name
        if any(name.startswith(p) for p in ["服务建议书", "服务验收表", "测试用例"]):
            print(f"  {name} ({f.stat().st_size/1024:.0f} KB)")
    print()


if __name__ == "__main__":
    main()
